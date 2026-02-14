from pydantic import BaseModel, Field
# from langchain_community.document_loaders import DirectoryLoader
from glob import glob
from pathlib import Path
import constants as const
from config import Config
import constants as const
import os
from prompts import SystemPrompts
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from pydantic_store import Chunk, Chunks
from langchain_core.output_parsers import PydanticOutputParser
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_chroma import Chroma
from dotenv import load_dotenv
import traceback
from docx import Document as DocxDocument
from langchain_core.documents import Document

load_dotenv(override=True)

llm = ChatOpenAI(
	model=const.MODEL_NAME,
	temperature=0
)
embedding = OpenAIEmbeddings(model="text-embedding-3-small")

# def _fetch_documents():
#     folders = glob(str(Path(Config.KNOWLEDGE_BASE_PATH) / "*"))
#     documents = []
#     for folder in folders:
#         doc_type = os.path.basename(folder)
#         loader = DirectoryLoader(
#             folder,
#             glob="**/*.docx",
#             loader_cls=UnstructuredWordDocumentLoader
#         )
#         folder_docs = loader.load()
#         for doc in folder_docs:
#             doc.metadata["doc_type"] = doc_type
#             documents.append(doc)
#     return documents

def load_docx_file(path: str) -> list[Document]:
    doc = DocxDocument(path)

    paragraphs = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip()
    ]

    text = "\n\n".join(paragraphs)

    return [
        Document(
            page_content=text,
            metadata={"source":path}
        )
    ]

def fetch_documents():
    status = False
    documents = []
    try:
        folders = glob(str(Path(Config.KNOWLEDGE_BASE_PATH) / "*"))
        
        for folder in folders:
            doc_type = os.path.basename(folder)
            files = glob(os.path.join(folder, "**/*.docx"), recursive=True)

            for file_path in files:
                docs = load_docx_file(file_path)

                for doc in docs:
                    doc.metadata["type"] = doc_type
                    documents.append(doc)
        status = True
    except Exception as err:
        print("Error in fetch_document :", err)
    finally:
        return status, documents

def process_document(documents):
    status = False
    all_chunks = []
    try:
        structured_llm = llm.with_structured_output(Chunks)
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", "Return ONLY valid JSON. No explanations."),
            ('human', SystemPrompts.INGEST_PROMPT)
        ])

        chain = prompt | structured_llm

        inputs = []

        for doc in documents:
            how_many = (len(doc.page_content) // const.AVERAGE_CHUNK_SIZE) + 1

            inputs.append({
                "doc_type": doc.metadata.get("doc_type"),
                "doc_source": doc.metadata.get("source"),
                "how_many": how_many,
                "doc_text": doc.page_content
            })

        print(f"Processing {len(inputs)} documents...")

        outputs: list[Chunks] = chain.batch(
            inputs,
            config={"max_concurrency": 1}
        )
        for source_doc, outputs in zip(documents, outputs):
            for chunk in outputs.chunks:
                all_chunks.append(chunk.as_result(source_doc))
        print(f"All Chunks {all_chunks}")
        print(f"Created {len(all_chunks)} chunks")
        status = True
    except Exception as err:
        print("Error occured in process document:", err)
    finally:
        return status, all_chunks

def create_embeddings(chunks):
    try:
        if os.path.exists(Config.DB_NAME):
            Chroma(persist_directory=Config.DB_NAME, embedding_function=embedding).delete_collection()
            print("Delete Old ")

        vectorstore = Chroma.from_documents(
            documents=chunks, embedding=embedding, persist_directory=Config.DB_NAME
        )
        collection = vectorstore._collection
        count = collection.count()

        sample_embedding = collection.get(limit=1, include=["embeddings"])["embeddings"][0]
        dimensions = len(sample_embedding)
        print(f"There are {count:,} vectors with {dimensions:,} dimensions in the vector store")
        print("Data is Ingested")
    except Exception as err:
        print("Exception in Creating Embeddings :", err)



if __name__ == "__main__":
    status, documents = fetch_documents()
    if status:
        status, chunks = process_document(documents)
        if status:
            create_embeddings(chunks)
